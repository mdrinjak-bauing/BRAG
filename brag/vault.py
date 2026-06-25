"""General vault file access — makes BRAG a single MCP for BOTH the search corpus
AND the user's note vault(s).

The built-in notebook/passage tools are locked to a separate `WissensWIKI/` tree
(see config) and therefore cannot reach a user's real working folders. These helpers
lift that limit: read / list /
search / write / append / edit ANY text file under the vault root(s), so no separate
filesystem or Obsidian MCP is needed.

Multi-root: besides the default vault (VAULT_DIR), additional NAMED roots can be
configured via VAULT_EXTRA_ROOTS ("alias=/abs/path,…"). A path targets one with an
"alias:rest" prefix (e.g. "alt:reports/x.md"); no prefix = the default vault. Extra
roots are pure FILE access — they are NOT part of the search corpus.

vault_extract is the explicit counterpart for BINARY documents (PDF/Word/Excel):
vault_read stays text-only (Markdown/txt/csv), vault_extract is used only on demand.

Safety:
- Everything is confined to a configured root; a path that resolves outside every
  root (or names an unknown alias) is rejected.
- READS are allowed anywhere under a root.
- WRITES in the DEFAULT vault are blocked in the configured write-protected folders
  (config.VAULT_WRITE_PROTECT) and any hidden dir, and never overwrite an existing
  file unless overwrite=True. Extra roots block only hidden dirs (the configured
  top-level protections apply to the default vault).
"""

from __future__ import annotations

import os
from pathlib import Path

from brag import config

# Write-protected / search-skipped top-level folders are configured per deployment
# (config.VAULT_WRITE_PROTECT / config.VAULT_SEARCH_SKIP); empty by default.
_TEXT_EXT = {".md", ".markdown", ".txt", ".csv"}
# Binary formats vault_extract turns into text (on explicit request only).
_EXTRACT_EXT = {".pdf", ".docx", ".xlsx"}
_EXTRACT_MAX_CHARS = 50_000  # cap so a huge document never floods the context


def _default_root() -> Path:
    return Path(config.VAULT).resolve()


def _roots() -> dict:
    """alias → resolved root. "" is the default (primary) vault; the rest come
    from config.EXTRA_VAULT_ROOTS (pure file access, not the corpus)."""
    roots = {"": _default_root()}
    for alias, p in getattr(config, "EXTRA_VAULT_ROOTS", {}).items():
        try:
            roots[alias] = Path(p).expanduser().resolve()
        except Exception:  # noqa: BLE001
            pass
    return roots


def _split_alias(path: str) -> tuple[str, str]:
    """Split an optional 'alias:rest' prefix where alias is a CONFIGURED extra
    root. No match → ("", path) (the default vault). Guards against eating things
    like 'http:' — only a known alias counts."""
    if ":" in path:
        alias, _, rest = path.partition(":")
        if alias in getattr(config, "EXTRA_VAULT_ROOTS", {}):
            return alias, rest
    return "", path


def _resolve(path: str):
    """Resolve a (possibly alias-prefixed) vault path. Returns (resolved_path,
    root, alias) or None if it escapes its root / names an unknown alias."""
    alias, rest = _split_alias(path)
    root = _roots().get(alias)
    if root is None:
        return None
    raw = Path(rest)
    p = (raw if raw.is_absolute() else root / raw).resolve()
    try:
        p.relative_to(root)
    except ValueError:
        return None
    return p, root, alias


def _display(p: Path, root: Path, alias: str) -> str:
    rel = p.relative_to(root).as_posix()
    return f"{alias}:{rel}" if alias else rel


def _write_protected(p: Path, root: Path, alias: str) -> bool:
    rel = p.relative_to(root)
    parts = rel.parts
    if not parts:
        return True
    # corpus/code protection applies to the DEFAULT vault only
    if alias == "" and parts[0] in config.VAULT_WRITE_PROTECT:
        return True
    return any(part.startswith(".") for part in parts)


_OUTSIDE = "Pfad liegt außerhalb der erlaubten Vault-Wurzeln (oder unbekannter alias:): "
_PROTECTED = ("Schreibgeschützt: '{path}' liegt in einem geschützten Bereich "
              "(VAULT_WRITE_PROTECT) oder einem versteckten Ordner — hier schreibt "
              "das Tool nicht.")


def vault_read(path: str) -> str:
    r = _resolve(path)
    if r is None:
        return _OUTSIDE + path
    p, root, alias = r
    if not p.exists() or not p.is_file():
        return f"Datei nicht gefunden: {path}"
    try:
        return p.read_text(encoding="utf-8")
    except Exception as e:  # noqa: BLE001
        return (f"Konnte '{path}' nicht als Text lesen ({e}). Für PDF/Word/Excel "
                f"vault_extract benutzen (nur auf ausdrückliche Aufforderung).")


def vault_list(subdir: str = "") -> str:
    if subdir:
        r = _resolve(subdir)
        if r is None:
            return _OUTSIDE + subdir
        d, root, alias = r
    else:
        d = root = _default_root()
        alias = ""
    if not d.exists() or not d.is_dir():
        return f"Ordner nicht gefunden: {subdir or '(Wurzel)'}"
    rows = []
    for c in sorted(d.iterdir(), key=lambda x: (x.is_file(), x.name.lower())):
        if c.name == ".DS_Store" or c.name.startswith("."):
            continue
        label = _display(c, root, alias)
        if c.is_dir():
            n = sum(1 for _ in c.iterdir())
            rows.append(f"📁 {label}/  ({n})")
        else:
            rows.append(f"📄 {label}")
    return "\n".join(rows) if rows else "(leer)"


def vault_search(query: str, content: bool = True, limit: int = 40,
                 root: str = "") -> str:
    """Find vault FILES by name and (optionally) by content. Default = the primary
    vault; pass root="fh" (a configured extra root) to search THERE instead. Not the
    literature corpus — that is what search() is for. In the primary vault, code/
    corpus/archive areas are skipped; extra roots are searched whole."""
    roots = _roots()
    base = roots.get(root)
    if base is None:
        avail = ", ".join(a or "(default)" for a in roots)
        return f"Unbekannte Vault-Wurzel '{root}'. Verfügbar: {avail}."
    q = query.lower().strip()
    if not q:
        return "Leere Suchanfrage."
    skip_top = config.VAULT_SEARCH_SKIP if root == "" else set()
    prefix = f"{root}:" if root else ""
    hits: list[tuple[str, str, str]] = []
    for dirpath, dirnames, filenames in os.walk(base):
        rel_parts = Path(dirpath).resolve().relative_to(base).parts
        # prune hidden dirs everywhere, and the skip-areas at the top level
        dirnames[:] = [
            d for d in dirnames
            if not d.startswith(".") and not (len(rel_parts) == 0 and d in skip_top)
        ]
        for fn in filenames:
            if fn == ".DS_Store" or fn.startswith("."):
                continue
            p = Path(dirpath) / fn
            rel = prefix + p.relative_to(base).as_posix()
            name_hit = q in fn.lower()
            snippet = ""
            content_hit = False
            if content and not name_hit and p.suffix.lower() in _TEXT_EXT:
                try:
                    txt = p.read_text(encoding="utf-8", errors="ignore")
                    i = txt.lower().find(q)
                    if i >= 0:
                        content_hit = True
                        s = max(0, i - 40)
                        snippet = txt[s:i + 80].replace("\n", " ").strip()
                except Exception:  # noqa: BLE001
                    pass
            if name_hit or content_hit:
                hits.append((rel, "Name" if name_hit else "Inhalt", snippet))
                if len(hits) >= limit:
                    break
        if len(hits) >= limit:
            break
    if not hits:
        return f"Keine Vault-Datei zu '{query}'{f' in {root!r}' if root else ''} gefunden."
    out = [f"**{len(hits)} Vault-Treffer** zu '{query}'"
           + (f" (Wurzel {root})" if root else "") + ":"]
    for rel, kind, snip in hits:
        out.append(f"- `{rel}` ({kind})" + (f" — …{snip}…" if snip else ""))
    return "\n".join(out)


def vault_write(path: str, content: str, overwrite: bool = False) -> str:
    r = _resolve(path)
    if r is None:
        return _OUTSIDE + path
    p, root, alias = r
    if _write_protected(p, root, alias):
        return _PROTECTED.format(path=path)
    if p.exists() and not overwrite:
        return (f"'{path}' existiert bereits. Zum Überschreiben overwrite=True setzen "
                f"(vorher abklären), vault_append zum Anhängen oder vault_edit "
                f"für eine gezielte Änderung nutzen.")
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"✓ geschrieben: {_display(p, root, alias)} ({len(content)} Zeichen)"
    except Exception as e:  # noqa: BLE001
        return f"Konnte '{path}' nicht schreiben: {e}"


def vault_append(path: str, content: str) -> str:
    r = _resolve(path)
    if r is None:
        return _OUTSIDE + path
    p, root, alias = r
    if _write_protected(p, root, alias):
        return _PROTECTED.format(path=path)
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        existed = p.exists()
        sep = "" if (not existed or p.read_text(encoding="utf-8").endswith("\n")) else "\n"
        with p.open("a", encoding="utf-8") as f:
            f.write(sep + content)
        verb = "angehängt an" if existed else "neu angelegt"
        return f"✓ {verb}: {_display(p, root, alias)} (+{len(content)} Zeichen)"
    except Exception as e:  # noqa: BLE001
        return f"Konnte an '{path}' nicht anhängen: {e}"


def vault_edit(path: str, old_string: str, new_string: str,
               replace_all: bool = False) -> str:
    """Chirurgischer In-Place-Edit: ersetzt einen exakten Textausschnitt in einer
    BESTEHENDEN Vault-Datei — das Gegenstück zu vault_write (ganze Datei) und
    vault_append (nur ans Ende). Damit lässt sich der Status oben in einer
    Konzeptnotiz auffrischen, eine einzelne Zeile im Themen-Katalog der Übersicht
    nachziehen oder ein falsch gespeicherter Beleg punktuell korrigieren, ohne die
    ganze Datei neu zu schreiben.

    Wie der Code-Editor: old_string muss EXAKT passen (inkl. Einrückung und
    Zeilenumbrüchen) und EINDEUTIG sein, sonst wird der Edit verweigert — mehr
    Kontext in old_string aufnehmen oder replace_all=True für alle Vorkommen.
    Gleicher Schreibschutz wie vault_write; legt keine neue Datei an (dafür
    vault_write)."""
    r = _resolve(path)
    if r is None:
        return _OUTSIDE + path
    p, root, alias = r
    if _write_protected(p, root, alias):
        return _PROTECTED.format(path=path)
    if not p.exists() or not p.is_file():
        return (f"Datei nicht gefunden: {path} — vault_edit ändert nur bestehende "
                f"Dateien; zum Neuanlegen vault_write nutzen.")
    if old_string == "":
        return "old_string ist leer — gib den exakt zu ersetzenden Text an."
    if old_string == new_string:
        return "old_string und new_string sind identisch — nichts zu ändern."
    try:
        text = p.read_text(encoding="utf-8")
    except Exception as e:  # noqa: BLE001
        return f"Konnte '{path}' nicht als Text lesen: {e}"
    count = text.count(old_string)
    if count == 0:
        return (f"Text nicht gefunden in '{path}'. old_string muss EXAKT passen "
                f"(inkl. Einrückung/Zeilenumbrüche) — vorher mit vault_read den "
                f"genauen Wortlaut prüfen.")
    if count > 1 and not replace_all:
        return (f"old_string kommt {count}× in '{path}' vor — nicht eindeutig. Mehr "
                f"Kontext aufnehmen, damit es genau einmal passt, oder replace_all=True "
                f"setzen, um alle {count} Vorkommen zu ersetzen.")
    new_text = text.replace(old_string, new_string)
    try:
        p.write_text(new_text, encoding="utf-8")
    except Exception as e:  # noqa: BLE001
        return f"Konnte '{path}' nicht schreiben: {e}"
    delta = len(new_text) - len(text)
    sign = "+" if delta >= 0 else ""
    verb = f"{count}× ersetzt" if replace_all and count > 1 else "ersetzt"
    return f"✓ {verb} in {_display(p, root, alias)} ({sign}{delta} Zeichen)"


# ── vault_extract: Text aus Binär-Dokumenten (PDF/Word/Excel) — nur auf Wunsch ──

class _MissingDep(Exception):
    """Eine optionale Extraktions-Bibliothek fehlt — klare pip-Meldung statt Crash."""


def _extract_pdf(p: Path, page_from: int, page_to: int) -> str:
    try:
        import pypdfium2 as pdfium
    except Exception as e:  # noqa: BLE001
        raise _MissingDep("PDF-Extraktion braucht pypdfium2 — `pip install pypdfium2`.") from e
    doc = pdfium.PdfDocument(str(p))
    try:
        n = len(doc)
        start = (max(1, page_from) - 1) if page_from else 0
        end = min(n, page_to) if page_to else n
        parts = []
        for i in range(start, max(start, end)):
            page = doc[i]
            tp = page.get_textpage()
            parts.append(f"[S. {i + 1}]\n" + tp.get_text_range())
            tp.close()
            page.close()
        return "\n\n".join(parts)
    finally:
        doc.close()


def _extract_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except Exception as e:  # noqa: BLE001
        raise _MissingDep("Word-Extraktion braucht python-docx — `pip install python-docx`.") from e
    d = docx.Document(str(p))
    lines = [para.text for para in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            lines.append(" | ".join(c.text for c in row.cells))
    return "\n".join(line for line in lines if line)


def _extract_xlsx(p: Path) -> str:
    try:
        import openpyxl
    except Exception as e:  # noqa: BLE001
        raise _MissingDep("Excel-Extraktion braucht openpyxl — `pip install openpyxl`.") from e
    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    try:
        out = []
        for ws in wb.worksheets:
            out.append(f"## Blatt: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    out.append(" | ".join(cells))
        return "\n".join(out)
    finally:
        wb.close()


def vault_extract(path: str, page_from: int = 0, page_to: int = 0) -> str:
    """Extrahiert TEXT aus einer Vault-PDF/-Word/-Excel — **nur auf ausdrückliche
    Aufforderung** (Standard ist vault_read für Markdown/Text). Erkennt das Format
    an der Endung (.pdf/.docx/.xlsx); read-only; auf die erlaubten Wurzeln begrenzt
    (auch fh:). PDF optional auf einen Seitenbereich (page_from/page_to, 1-basiert)
    einschränken. Lange Dokumente werden gedeckelt."""
    r = _resolve(path)
    if r is None:
        return _OUTSIDE + path
    p, root, alias = r
    if not p.exists() or not p.is_file():
        return f"Datei nicht gefunden: {path}"
    ext = p.suffix.lower()
    if ext not in _EXTRACT_EXT:
        if ext in _TEXT_EXT:
            return (f"'{path}' ist bereits Text — dafür vault_read benutzen "
                    f"(vault_extract ist für PDF/Word/Excel).")
        return (f"vault_extract unterstützt {sorted(_EXTRACT_EXT)} — nicht '{ext or '?'}'.")
    try:
        if ext == ".pdf":
            text = _extract_pdf(p, page_from, page_to)
        elif ext == ".docx":
            text = _extract_docx(p)
        else:  # .xlsx
            text = _extract_xlsx(p)
    except _MissingDep as e:
        return str(e)
    except Exception as e:  # noqa: BLE001
        return f"Konnte '{path}' nicht extrahieren: {e}"
    text = text or "(kein extrahierbarer Text — evtl. ein Scan ohne Textebene?)"
    truncated = len(text) > _EXTRACT_MAX_CHARS
    body = text[:_EXTRACT_MAX_CHARS]
    head = (f"📄 Extrahiert aus {_display(p, root, alias)} "
            f"({ext[1:]}, {len(text)} Zeichen{', gekürzt' if truncated else ''}):\n\n")
    tail = "\n\n[gekürzt — mit page_from/page_to einschränken]" if truncated else ""
    return head + body + tail
